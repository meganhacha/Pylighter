from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from tkinter import Tk, filedialog
import os

COLOR_INDEX = {
   "PINK": WD_COLOR_INDEX.PINK,
   "GREEN": WD_COLOR_INDEX.GREEN,
   "TURQUOISE": WD_COLOR_INDEX.TURQUOISE,
   "YELLOW": WD_COLOR_INDEX.YELLOW,
   "BRIGHT_GREEN": WD_COLOR_INDEX.BRIGHT_GREEN,
}

underline_list = []

#Bold event titles and subsequent info if on the same line.
def bold_text(doc, keyword):

   for p in doc.paragraphs:
      for r in p.runs:
         if keyword in r.text:
            r.font.bold = True

#Underline dates
def underline_text(doc, keyword):
   for p in doc.paragraphs:
      
      for r in p.runs:
         if any(word in r.text for word in keyword):
            r.font.underline = True

#Change the font and text size of the entire document.
def allover_text(doc):
   for p in doc.paragraphs:

      for r in p.runs:
         r.font.size = Pt(12)
         r.font.name = 'Arial'

#Get each individual line, ignoring whitespaces.
def get_paragraph_text(paragraph):
   return ''.join(run.text for run in paragraph.runs).strip()


#Highlight the lines that are applicable, based on the list of keywords below.
def highlight_lines(doc, keyword, color):

   for paragraph in doc.paragraphs:
      if keyword in get_paragraph_text(paragraph):

         for run_h in paragraph.runs:
            run_h.font.highlight_color = color
            if keyword in {"Event:", "Meeting:", "SETUP"}:
               run_h.font.color.rgb = RGBColor(255, 255, 255)
            

#Allows the saving of a new -- now highlighted and edited -- Word doc.
def save_doc(doc, output_path):
   doc.save(output_path)

#Asks the user to select the .docx file they would like to perform the highlighting and other edits on.
def select_file():
   root = Tk()
   root.withdraw()

   file_path = filedialog.askopenfilename(
      title = "Select docx File",
      filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
   )

   return file_path

#Using the provided keywords, selects the highlight color for each section.
def get_highlight(keyword):
   key_color_mapping = {
        "PINK": {"MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"},
        "GREEN": {"Event:", "Meeting:", "SETUP"},
        "TURQUOISE": {"Location:"},
        "YELLOW": {"Event Date/Time:", "Date/Time:"},
        "BRIGHT_GREEN": {"Concierge:", "Concierge/MSGS:", "Concierge/FOH:"},
    }

   for color, keys in key_color_mapping.items():
      if keyword in keys:
         return COLOR_INDEX[color]

   return None


def main():
   
    input_file = select_file()

    if not input_file:
      print("No file selected. Exiting")
    else:
      #Open the selected doc, and add the applicable keywords that need to have their sections highlighted.
      doc= Document(input_file)
      highlight_list = {"Event:", "Meeting:", "SETUP", "Event Date/Time:", "Date/Time:", "Location:", "MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY", "Concierge:", "Concierge/MSGS:", "Concierge/FOH:"}
      bold_list = {"Organization:", "Contact person day of event:", "# of people expected to attend:", "Event Description:", "Setup/Tear Down:", "Tech Needs:", "Food Services:", "Security"}

      #Change the text to the correct font and text size.
      allover_text(doc)

      #Use the bold list to bold all instances of the bold list in the Document.
      for b_word in bold_list:
         bold_text(doc, b_word)
      for h_word in highlight_list:
         bold_text(doc, h_word)

      #Traverse the document, highlighting each word with its corresponding color.
      for key in highlight_list:
         color = get_highlight(key)
         highlight_lines(doc, key, color)
      
      

      output_file = filedialog.asksaveasfilename(
         defaultextension=".docx",
         filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]

       )

      if output_file:
         save_doc(doc, output_file)
         print(f"Document saved to {output_file}")
      else:
         print("No output file selected. Exiting.")       


if __name__ == "__main__":
    main()
